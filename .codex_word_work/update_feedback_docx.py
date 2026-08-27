from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image


BASE = Path(
    r"C:\Users\HOSEO\OneDrive - 호서대학교\나는 개인이요\석사생\논문"
    r"\Kalman Filter\논문\latex양식\Elsevier"
)
SOURCE = BASE / "논문_한글판_피드백용_v09.docx"
FIGURES = BASE / "figures" / "font_size_candidates"
OUTPUT = Path(__file__).resolve().parent / "논문_한글판_피드백용_v10.docx"


IMAGE_SPECS = [
    ("sensor_array.png", 3.55),
    ("transformer_blockdiagram.png", 3.05),
    ("trajectory_xy_plane.png", 4.75),
    ("trajectory_xz_plane.png", 4.75),
    ("trajectory_yz_plane.png", 3.80),
    ("trajectory_3d.png", 3.80),
    ("rmse_distance_0_600m.png", 5.35),
    ("rmse_tdoa_bias_0_100us.png", 3.00),
    ("rmse_doa_input_angular_error_std_0_1p2deg.png", 3.00),
    ("rmse_tdoa_random_input_error_std_0_100us.png", 4.35),
]


def iter_inline_shapes(document: Document):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for drawing in run._r.xpath(".//w:drawing"):
                inline = drawing.find(qn("wp:inline"))
                if inline is not None:
                    yield inline


def replace_images(document: Document) -> None:
    shapes = list(iter_inline_shapes(document))
    if len(shapes) != len(IMAGE_SPECS):
        raise RuntimeError(
            f"Expected {len(IMAGE_SPECS)} inline images, found {len(shapes)}"
        )

    for inline, (filename, width_inches) in zip(shapes, IMAGE_SPECS):
        image_path = FIGURES / filename
        with Image.open(image_path) as image:
            pixel_width, pixel_height = image.size

        blip = inline.find(".//" + qn("a:blip"))
        if blip is None:
            raise RuntimeError(f"Missing image relationship for {filename}")
        relation_id = blip.get(qn("r:embed"))
        image_part = document.part.related_parts[relation_id]
        image_part._blob = image_path.read_bytes()

        width = Inches(width_inches)
        height = Inches(width_inches * pixel_height / pixel_width)
        extent = inline.find(qn("wp:extent"))
        transform_extent = inline.find(".//" + qn("a:xfrm") + "/" + qn("a:ext"))
        if extent is None or transform_extent is None:
            raise RuntimeError(f"Missing drawing extent for {filename}")
        for node in (extent, transform_extent):
            node.set("cx", str(width))
            node.set("cy", str(height))


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr

    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths_dxa)))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = width
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(width))


def update_text(document: Document) -> None:
    old_caption = (
        "Fig. 3 (continued). Trajectory estimation comparison (continued) in "
        "(c) Y-Z and (d) 3D views. Styles: ground truth, black dashed; proposed, "
        "orange solid/circles; LSTM, magenta dash-dot/squares; MLP, blue "
        "dotted/triangles; 1D-CNN, cyan long-dashed/crosses; MUSIC-Hybrid, "
        "green dashed/diamonds."
    )
    new_caption = (
        "Fig. 3 (continued). Trajectory estimation comparison in (c) Y-Z and "
        "(d) 3D views. Line and marker styles are defined in Fig. 3(a)-(b)."
    )
    matches = [p for p in document.paragraphs if p.text == old_caption]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one continued Fig. 3 caption, found {len(matches)}")
    matches[0].text = new_caption
    matches[0].style = document.styles["Caption"]


def main() -> None:
    document = Document(SOURCE)
    replace_images(document)
    set_table_geometry(document.tables[0], [5616, 3744])
    update_text(document)
    document.core_properties.modified = datetime.now()
    document.core_properties.title = (
        "Transformer-Based Feature-Level TOA/TDOA/DOA Fusion for Compact USBL Localization"
    )
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
